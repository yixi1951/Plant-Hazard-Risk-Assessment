# -*- coding: utf-8 -*-
"""
生成论文中文 Word 模板
基于 python-docx，包含封面、摘要、正文、参考文献等完整结构。
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# ── 输出路径 ──────────────────────────────────────────────────
OUTPUT_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "论文模板_农作物病害智能诊断系统.docx")


# ====================================================================
# 工具函数
# ====================================================================

def set_cell_shading(cell, color_hex):
    """设置表格单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, font_name_cn="宋体", font_name_en="Times New Roman", size=12, bold=False, italic=False, color=None):
    """统一设置 run 的字体属性"""
    run.font.size = Pt(size)
    run.font.name = font_name_en
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text, font_cn="宋体", font_en="Times New Roman", size=12,
                  bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=0, line_spacing=1.5,
                  first_line_indent=None, italic=False, color=None,
                  keep_with_next=False):
    """添加一个格式化的段落"""
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(first_line_indent)
    if keep_with_next:
        para.paragraph_format.keep_with_next = True

    run = para.add_run(text)
    set_run_font(run, font_cn, font_en, size, bold, italic, color)
    return para


def add_heading_custom(doc, text, level=1):
    """添加自定义格式的标题（符合中文论文规范）"""
    fonts = {0: ("黑体", "Arial", 22, True),   # 论文题目
             1: ("黑体", "Arial", 16, True),   # 一级标题（三号 ≈ 16pt）
             2: ("黑体", "Arial", 14, True),   # 二级标题（四号 ≈ 14pt）
             3: ("黑体", "Arial", 12, True)}   # 三级标题（小四 ≈ 12pt）
    font_cn, font_en, size, bold = fonts.get(level, ("黑体", "Arial", 12, True))

    if level == 0:
        alignment = WD_ALIGN_PARAGRAPH.CENTER
        space_before, space_after = 24, 12
    else:
        alignment = WD_ALIGN_PARAGRAPH.LEFT
        space_before, space_after = 18, 6

    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.keep_with_next = True

    run = para.add_run(text)
    set_run_font(run, font_cn, font_en, size, bold)
    return para


def add_body_text(doc, text, first_line_indent=0.74):
    """添加正文段落（宋体 12pt，首行缩进 2 字符）"""
    return add_paragraph(
        doc, text, font_cn="宋体", size=12, bold=False,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=1.5, first_line_indent=first_line_indent,
        space_before=0, space_after=0
    )


def add_table_with_style(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        set_run_font(run, "黑体", "Arial", 10, bold=True, color=(255, 255, 255))
        set_cell_shading(cell, "4472C4")

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(cell_text))
            set_run_font(run, "宋体", "Times New Roman", 10)
            # 交替行底色
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D9E2F3")

    # 列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_figure_placeholder(doc, caption, width_cm=12, height_lines=6):
    """添加图片占位框"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)

    # 创建一个带底色的段落模拟图片框
    run = para.add_run(f"  [ 此处插入：{caption} ]  ")
    set_run_font(run, "宋体", "Arial", 11, italic=True, color=(100, 100, 100))

    # 图注
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.paragraph_format.space_before = Pt(3)
    caption_para.paragraph_format.space_after = Pt(6)
    caption_para.paragraph_format.line_spacing = 1.5
    run = caption_para.add_run(f"图 1  {caption}")
    set_run_font(run, "宋体", "Times New Roman", 10, bold=False)
    return para


def add_table_caption(doc, caption):
    """添加表格标题"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(f"表 1  {caption}")
    set_run_font(run, "黑体", "Times New Roman", 10, bold=True)


def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()


def setup_page(doc):
    """设置页面布局"""
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)


def add_footer_page_number(doc):
    """添加页码（居中）"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 使用 Word 页码域
        run = para.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        run2 = para.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)
        run3 = para.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)


# ====================================================================
# 生成论文模板
# ====================================================================

def generate_template():
    doc = Document()
    setup_page(doc)

    # ──────────────────────────────────────────────────────────────
    # 封面页
    # ──────────────────────────────────────────────────────────────
    # 空行
    for _ in range(4):
        add_paragraph(doc, "", size=12, line_spacing=1.5)

    # 学校名称
    add_paragraph(doc, "XXXX 大学", font_cn="黑体", size=26, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # 论文类型
    add_paragraph(doc, "本科毕业论文 / 硕士研究生学位论文", font_cn="黑体", size=16, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # 论文题目
    add_paragraph(doc, "基于多任务深度学习的农作物病害", font_cn="黑体", size=22, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=0)
    add_paragraph(doc, "智能诊断与风险评估系统", font_cn="黑体", size=22, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=36)

    for _ in range(3):
        add_paragraph(doc, "", size=12, line_spacing=1.5)

    # 封面信息表
    cover_info = [
        ("学    院：", "XXXX 学院"),
        ("专    业：", "计算机科学与技术 / 人工智能"),
        ("学    号：", "XXXXXXXXXX"),
        ("学生姓名：", "XXX"),
        ("指导教师：", "XXX 教授"),
    ]
    for label, value in cover_info:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = 2.0
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run1 = para.add_run(label)
        set_run_font(run1, "宋体", "Times New Roman", 14, bold=True)
        run2 = para.add_run(value)
        set_run_font(run2, "宋体", "Times New Roman", 14, bold=False)
        run2.font.underline = True

    # 日期
    for _ in range(2):
        add_paragraph(doc, "", size=12, line_spacing=1.5)
    add_paragraph(doc, "二〇二六年六月", font_cn="宋体", size=14, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_page_break(doc)

    # ──────────────────────────────────────────────────────────────
    # 中文摘要
    # ──────────────────────────────────────────────────────────────
    add_heading_custom(doc, "摘  要", level=1)

    abstract_text = (
        "农作物病害是威胁全球粮食安全的主要因素之一，快速准确的病害诊断对农业生产具有重要的现实意义。"
        "现有深度学习方法在植物病害图像识别领域取得了显著进展，但大多聚焦于单一任务（如仅病害分类或仅严重程度分级），"
        "缺乏对病害种类、严重程度及风险的联合诊断能力。"
    )
    abstract_text += (
        "本文提出了一种基于多任务深度学习的农作物病害智能诊断与风险评估框架。"
        "该框架以 MobileNetV2 为共享骨干网络，引入交叉注意力机制（Cross-Task Attention）为病害分类和严重程度分级"
        "分别生成任务特定特征，实现了病害类别识别（61 类）、严重程度分级（3 级）和风险评分的联合预测。"
        "同时，结合 Grad-CAM 可解释性方法生成热力图，定位图像中的病灶区域，提升模型决策的可解释性。"
    )
    abstract_text += (
        "在 PlantVillage 和 AgriculturalDisease 竞赛数据集上的实验结果表明："
        "多任务模型在疾病分类准确率和严重程度分级 F1 分数上均优于对应的单任务基线模型，"
        "验证了多任务学习在农业病害诊断任务中的协同效应。"
        "此外，本文还构建了完整的 Web 诊断系统，支持图片上传、实时推理、诊断报告生成等功能，"
        "具备良好的实际应用价值。"
    )

    # 关键词
    add_body_text(doc, abstract_text)

    # 关键词段落
    kw_para = doc.add_paragraph()
    kw_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw_para.paragraph_format.line_spacing = 1.5
    kw_para.paragraph_format.space_before = Pt(12)
    kw_para.paragraph_format.first_line_indent = Cm(0.74)
    run_label = kw_para.add_run("关键词：")
    set_run_font(run_label, "黑体", "Arial", 12, bold=True)
    run_kw = kw_para.add_run("多任务学习；农作物病害诊断；MobileNetV2；交叉注意力机制；Grad-CAM；风险评估")
    set_run_font(run_kw, "宋体", "Times New Roman", 12)

    add_page_break(doc)

    # ──────────────────────────────────────────────────────────────
    # 英文摘要
    # ──────────────────────────────────────────────────────────────
    add_heading_custom(doc, "Abstract", level=1)

    en_abstract = (
        "Crop diseases are one of the major threats to global food security. "
        "Rapid and accurate disease diagnosis is of great practical significance for agricultural production. "
        "Existing deep learning methods have achieved remarkable progress in plant disease image recognition, "
        "yet most focus on single tasks such as disease classification or severity grading alone, "
        "lacking the capability for joint diagnosis of disease type, severity level, and risk assessment."
    )
    en_abstract += (
        "This paper proposes a multi-task deep learning framework for intelligent crop disease diagnosis and risk assessment. "
        "The framework employs MobileNetV2 as a shared backbone network and introduces a Cross-Task Attention mechanism "
        "to generate task-specific features for disease classification and severity grading respectively, "
        "enabling joint prediction of disease category (61 classes), severity level (3 levels), and risk score. "
        "Furthermore, Grad-CAM is integrated to generate heatmaps that localize lesion regions, "
        "enhancing the interpretability of model decisions."
    )
    en_abstract += (
        "Experimental results on the PlantVillage and AgriculturalDisease benchmark datasets demonstrate that "
        "the multi-task model outperforms its single-task counterparts in both disease classification accuracy "
        "and severity grading F1 score, validating the synergy effect of multi-task learning "
        "in agricultural disease diagnosis tasks. "
        "A complete web-based diagnosis system is also developed, supporting image upload, real-time inference, "
        "and diagnostic report generation, demonstrating strong practical application value."
    )
    add_body_text(doc, en_abstract)

    # Keywords
    kw_para_en = doc.add_paragraph()
    kw_para_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw_para_en.paragraph_format.line_spacing = 1.5
    kw_para_en.paragraph_format.space_before = Pt(12)
    kw_para_en.paragraph_format.first_line_indent = Cm(0.74)
    run_label_en = kw_para_en.add_run("Keywords: ")
    set_run_font(run_label_en, "Arial", "Arial", 12, bold=True)
    run_kw_en = kw_para_en.add_run(
        "Multi-task Learning; Crop Disease Diagnosis; MobileNetV2; Cross-Task Attention; Grad-CAM; Risk Assessment"
    )
    set_run_font(run_kw_en, "Times New Roman", "Times New Roman", 12)

    add_page_break(doc)

    # ──────────────────────────────────────────────────────────────
    # 目录页（占位）
    # ──────────────────────────────────────────────────────────────
    add_heading_custom(doc, "目  录", level=1)
    add_paragraph(doc, "（请在 Word 中插入自动目录：引用 → 目录 → 自动目录）",
                  font_cn="宋体", size=12, color=(150, 150, 150),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24)

    # 手动目录示意
    toc_items = [
        ("摘要", 1), ("Abstract", 1),
        ("第 1 章  引言", 1),
        ("  1.1  研究背景与意义", 2),
        ("  1.2  国内外研究现状", 2),
        ("  1.3  本文主要贡献", 2),
        ("  1.4  本文组织结构", 2),
        ("第 2 章  相关工作", 1),
        ("  2.1  深度学习在植物病害识别中的应用", 2),
        ("  2.2  多任务学习", 2),
        ("  2.3  可解释人工智能", 2),
        ("第 3 章  基于多任务学习的病害诊断方法", 1),
        ("  3.1  问题定义与总体框架", 2),
        ("  3.2  多任务共享骨干网络", 2),
        ("  3.3  交叉注意力机制", 2),
        ("  3.4  任务特定分类头", 2),
        ("  3.5  损失函数设计", 2),
        ("  3.6  Grad-CAM 可解释性分析", 2),
        ("  3.7  风险等级评估方法", 2),
        ("第 4 章  实验设计与结果分析", 1),
        ("  4.1  数据集与预处理", 2),
        ("  4.2  实验设置与评价指标", 2),
        ("  4.3  多任务与单任务对比实验", 2),
        ("  4.4  消融实验", 2),
        ("  4.5  可视化分析", 2),
        ("  4.6  少样本学习实验", 2),
        ("第 5 章  讨论", 1),
        ("第 6 章  结论与展望", 1),
        ("参考文献", 1),
        ("致谢", 1),
        ("附录 A  主要代码片段", 1),
        ("附录 B  诊断报告样例", 1),
    ]
    for item, level in toc_items:
        indent = Cm(0) if level == 1 else Cm(0.8)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.line_spacing = 1.8
        para.paragraph_format.first_line_indent = indent
        run = para.add_run(item)
        size = 12 if level == 1 else 11
        set_run_font(run, "宋体", "Times New Roman", size, bold=(level == 1))

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 第 1 章  引言
    # ══════════════════════════════════════════════════════════════
    add_heading_custom(doc, "第 1 章  引言", level=1)

    add_heading_custom(doc, "1.1  研究背景与意义", level=2)
    add_body_text(doc,
        "农作物病害是制约全球农业生产的主要因素之一。据联合国粮食及农业组织（FAO）统计，"
        "全球每年因植物病害造成的粮食损失高达 10%~16%，对粮食安全和农业可持续发展构成严重威胁。"
        "传统的农作物病害诊断主要依赖农业技术人员的田间观察和实验室检测，存在效率低、主观性强、"
        "专业门槛高等问题，难以满足大规模、实时化的现代农业需求。"
    )
    add_body_text(doc,
        "近年来，以深度学习为代表的计算机视觉技术在植物病害图像识别领域取得了显著进展。"
        "从早期的卷积神经网络（CNN）到 ResNet、MobileNet、EfficientNet 等先进架构，"
        "图像分类的准确率不断提升。然而，现有工作大多聚焦于单一的分类任务，"
        "即仅判断病害种类或仅评估严重程度，缺乏对病害信息的全面联合诊断能力。"
    )
    add_body_text(doc,
        "针对上述问题，本文提出了一种基于多任务深度学习的农作物病害智能诊断与风险评估框架。"
        "该框架通过共享骨干网络和交叉注意力机制，同时完成病害分类、严重程度分级和风险评分三个任务，"
        "旨在为农业生产提供更全面、更精准的智能化诊断方案。"
    )

    add_heading_custom(doc, "1.2  国内外研究现状", level=2)
    add_body_text(doc,
        "在植物病害图像识别方面，Mohanty et al.（2016）使用 AlexNet 和 GoogLeNet 在 PlantVillage 数据集上"
        "实现了 99.35% 的分类准确率，开创了深度学习应用于植物病害识别的先河。"
        "Too et al.（2019）系统比较了 VGG16、ResNet50、InceptionV3、DenseNet121 等深度模型在植物病害分类中的表现。"
        "近年来，MobileNet 等轻量级网络因其参数少、推理快的特点，特别适合边缘端和移动端的部署。"
    )
    add_body_text(doc,
        "在多任务学习方面，Caruana（1997）提出了多任务学习的通用框架，"
        "通过共享表示实现多个相关任务的联合学习。"
        "Misra et al.（2016）提出 Cross-stitch Networks 在任务间共享和交换特征。"
        "Liu et al.（2019）提出 MTAN（Multi-Task Attention Network）引入注意力机制进行任务特定特征提取。"
        "然而，多任务学习在农业病害诊断领域的应用研究尚不充分。"
    )
    add_body_text(doc,
        "在可解释人工智能方面，Selvaraju et al.（2017）提出的 Grad-CAM 方法通过梯度加权"
        "可视化卷积神经网络的关注区域，已成为最广泛使用的模型解释工具之一。"
        "在农业领域，Grad-CAM 可用于定位叶片病灶区域，提升诊断结果的可信度。"
    )

    add_heading_custom(doc, "1.3  本文主要贡献", level=2)
    contributions = [
        "提出了一个基于 MobileNetV2 的多任务学习框架，实现病害分类、严重程度分级和风险评分的联合预测。",
        "设计了一种交叉注意力机制（Cross-Task Attention），使不同任务能从共享特征中自适应提取任务相关信息。",
        "通过多任务与单任务的对比实验，定量验证了多任务学习在农业病害诊断中的协同效应。",
        "集成了 Grad-CAM 可解释性分析和结构化诊断报告生成，提升了系统的透明度和实用性。",
        "构建了完整的 Web 诊断系统，支持图片上传、实时推理和批量预测。",
    ]
    for i, c in enumerate(contributions, 1):
        add_body_text(doc, f"（{i}）{c}")

    add_heading_custom(doc, "1.4  本文组织结构", level=2)
    add_body_text(doc,
        "本文共分为六章。第一章为引言，介绍研究背景、相关工作及本文贡献。"
        "第二章回顾相关技术，包括深度学习病害识别、多任务学习和可解释 AI。"
        "第三章详细介绍本文提出的多任务诊断方法，包括网络架构、交叉注意力机制和损失函数。"
        "第四章报告实验设置和多组对比实验结果。第五章讨论研究发现的启示和局限性。"
        "第六章总结全文并展望未来研究方向。"
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 第 2 章  相关工作
    # ══════════════════════════════════════════════════════════════
    add_heading_custom(doc, "第 2 章  相关工作", level=1)

    add_heading_custom(doc, "2.1  深度学习在植物病害识别中的应用", level=2)
    add_body_text(doc,
        "（此处请扩展文献综述，系统性回顾基于深度学习的植物病害图像识别方法，"
        "包括传统 CNN→VGG→ResNet→MobileNet→EfficientNet→Vision Transformer 的发展脉络，"
        "以及在 PlantVillage、AI Challenger 等数据集上的代表性工作。）"
    )

    add_heading_custom(doc, "2.2  多任务学习", level=2)
    add_body_text(doc,
        "（此处请综述多任务学习的相关研究：Hard Parameter Sharing、"
        "Cross-stitch Networks、MMoE、MTAN 等方法，并讨论多任务学习在农业领域的应用现状。）"
    )

    add_heading_custom(doc, "2.3  可解释人工智能", level=2)
    add_body_text(doc,
        "（此处请综述可解释 AI 方法，特别是 Grad-CAM、Grad-CAM++、LayerCAM 等可视化方法，"
        "以及它们在农业和医学图像分析中的应用。）"
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 第 3 章  基于多任务学习的病害诊断方法（核心）
    # ══════════════════════════════════════════════════════════════
    add_heading_custom(doc, "第 3 章  基于多任务学习的病害诊断方法", level=1)

    add_heading_custom(doc, "3.1  问题定义与总体框架", level=2)
    add_body_text(doc,
        "本文将农作物病害诊断定义为三个联合任务：（1）病害分类——识别叶片图像所属的病害类别（共 61 类）；"
        "（2）严重程度分级——评估病害严重程度（0：健康，1：一般疾病，2：严重疾病）；"
        "（3）风险评估——综合病害种类和严重程度计算风险评分。"
        "总体框架如图 1 所示：输入叶片图像 → MobileNetV2 骨干网 → 交叉注意力模块 → "
        "病害分类头 + 严重程度分类头 → 病害标签 + 严重程度 + 风险评分。"
    )
    add_figure_placeholder(doc, "多任务诊断框架整体架构图")

    add_heading_custom(doc, "3.2  多任务共享骨干网络", level=2)
    add_body_text(doc,
        "本文选用 MobileNetV2（Sandler et al., 2018）作为共享骨干网络。"
        "MobileNetV2 采用倒残差结构（Inverted Residuals）和线性瓶颈（Linear Bottleneck），"
        "在保持较高分类精度的同时，参数量和计算量远小于 VGG 和 ResNet 等传统网络，"
        "有利于后续的移动端和边缘端部署。"
    )
    add_body_text(doc,
        "为平衡迁移学习和任务特定训练，本文冻结了 MobileNetV2 的大部分层，仅训练最后 3 层。"
        "骨干网络输出维度为 1280，经全局平均池化（Global Average Pooling）后得到一个 1280 维的共享特征向量。"
        "该共享特征随后输入交叉注意力模块，由不同任务头提取各自关注的子特征。"
    )

    add_heading_custom(doc, "3.3  交叉注意力机制", level=2)
    add_body_text(doc,
        "交叉注意力机制（Cross-Task Attention）是本文的核心创新点。"
        "其设计思想为：病害分类任务应关注病斑的形态和纹理特征，而严重程度分级任务则应关注病灶面积占比和扩散程度。"
        "两个任务虽然共享底层特征，但对特征各维度的关注度不同。"
    )
    add_body_text(doc,
        "具体而言，给定共享特征向量 x∈R^1280，交叉注意力机制的计算过程如下："
    )
    # 公式用文本
    add_body_text(doc,
        "Q_d = W_d · x,   Q_s = W_s · x,   V = W_v · x",
        first_line_indent=0.74
    )
    add_body_text(doc,
        "其中 W_d、W_s、W_v 分别为病害查询、严重程度查询和值映射的线性变换矩阵。"
        "然后计算注意力权重并生成任务特定特征："
    )
    add_body_text(doc,
        "A_d = softmax(Q_d · V),   f_d = A_d ⊙ V,   f_s = A_s ⊙ V",
        first_line_indent=0.74
    )
    add_body_text(doc,
        "最终得到的 f_d 和 f_s 分别输入病害分类头和严重程度分类头。"
        "这种设计使得两个任务可以自适应地从共享特征中选择所需的信息维度，"
        "既保留了多任务学习的参数共享优势，又提供了任务特定的特征提取能力。"
    )

    add_heading_custom(doc, "3.4  损失函数设计", level=2)
    add_body_text(doc,
        "本文采用加权多任务损失函数，将疾病分类损失和严重程度分级损失进行加权组合："
    )
    add_body_text(doc,
        "L_total = λ_1 · L_CE(y_d, ŷ_d) + λ_2 · L_CE(y_s, ŷ_s)",
        first_line_indent=0.74
    )
    add_body_text(doc,
        "其中 L_CE 为标准交叉熵损失，y_d、y_s 分别为病害标签和严重程度标签，"
        "λ_1、λ_2 为任务权重，默认设置为 λ_1=1.0、λ_2=0.8。"
        "严重程度任务的权重略低，是因为严重程度分级的类别数较少（3 类），学习相对简单。"
    )

    add_heading_custom(doc, "3.5  Grad-CAM 可解释性分析", level=2)
    add_body_text(doc,
        "为增强模型决策的可解释性，本文引入了 Grad-CAM（Gradient-weighted Class Activation Mapping）方法。"
        "Grad-CAM 利用目标类别对最后一层卷积特征图的梯度，计算每个特征通道的权重，"
        "从而生成与类别相关的热力图，叠加到原始图像上以可视化的方式展示模型做出决策的依据。"
    )
    add_figure_placeholder(doc, "Grad-CAM 热力图可视化示例")

    add_heading_custom(doc, "3.6  风险等级评估方法", level=2)
    add_body_text(doc,
        "基于病害类别和严重程度分级结果，本文进一步设计了三级风险评估规则："
    )
    risk_table_data = [
        ["高风险", "严重（2 级）且疾病置信度 > 80%", "需立即采取防控措施"],
        ["中风险", "一般（1 级）且疾病置信度 > 70%", "需密切监测并预防性施药"],
        ["低风险", "健康状态或置信度较低", "常规田间管理即可"],
        ["不确定", "疾病置信度 < 70%", "建议人工复核"],
    ]
    add_table_caption(doc, "风险等级划分标准")
    add_table_with_style(doc,
        ["风险等级", "判定条件", "应对策略"],
        risk_table_data,
        col_widths=[3.5, 5.5, 6.0]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 第 4 章  实验设计与结果分析
    # ══════════════════════════════════════════════════════════════
    add_heading_custom(doc, "第 4 章  实验设计与结果分析", level=1)

    add_heading_custom(doc, "4.1  数据集与预处理", level=2)
    add_body_text(doc,
        "本文使用 PlantVillage 公开数据集和 AgriculturalDisease 竞赛数据集进行实验。"
        "PlantVillage 数据集包含 14 种作物的 38 类病害图像，共约 54,000 张。"
        "AgriculturalDisease 数据集包含 10 种作物的 61 类病害图像，涵盖苹果、樱桃、玉米、"
        "葡萄、柑桔、桃、辣椒、马铃薯、草莓、番茄等主要经济作物。"
    )

    add_table_caption(doc, "数据集统计信息")
    add_table_with_style(doc,
        ["数据集", "作物种类", "病害类别数", "总样本量", "训练/验证比例"],
        [
            ["PlantVillage", "14", "38", "~54,000", "80% / 20%"],
            ["AgriculturalDisease", "10", "61", "~45,000", "80% / 20%"],
        ],
        col_widths=[3.5, 2.5, 2.5, 2.5, 2.5]
    )

    add_heading_custom(doc, "4.2  实验设置与评价指标", level=2)
    add_body_text(doc,
        "所有实验在以下统一条件下进行。"
    )

    add_table_caption(doc, "实验超参数设置")
    add_table_with_style(doc,
        ["超参数", "取值"],
        [
            ["骨干网络", "MobileNetV2（ImageNet 预训练）"],
            ["输入尺寸", "128 × 128"],
            ["优化器", "AdamW（lr=5e-4, weight_decay=1e-4）"],
            ["学习率调度", "ReduceLROnPlateau（patience=3, factor=0.5）"],
            ["早停轮次", "5"],
            ["最大训练轮次", "15"],
            ["批大小", "16（GPU）/ 4（CPU）"],
            ["数据增强", "RandomHorizontalFlip(p=0.3), RandomRotation(10°)"],
            ["任务权重", "λ₁=1.0（疾病）, λ₂=0.8（严重程度）"],
        ],
        col_widths=[5.0, 10.0]
    )

    add_heading_custom(doc, "4.3  多任务与单任务对比实验", level=2)
    add_body_text(doc,
        "为验证多任务学习的协同效应，本文训练了三组模型进行对比："
        "（1）多任务模型（本文方法），同时预测病害类别和严重程度；"
        "（2）单任务疾病分类模型，仅预测病害类别；"
        "（3）单任务严重程度分级模型，仅预测严重程度等级。"
        "三组模型使用相同的骨干网络和训练条件，确保对比公平。"
    )

    add_table_caption(doc, "主实验结果：多任务 vs 单任务性能对比")
    add_table_with_style(doc,
        ["模型", "疾病分类 Acc (%)", "疾病分类 F1", "严重程度 Acc (%)", "严重程度 F1"],
        [
            ["单任务疾病模型", "XX.XX", "X.XXXX", "—", "—"],
            ["单任务严重程度模型", "—", "—", "XX.XX", "X.XXXX"],
            ["多任务模型（本文）", "XX.XX", "X.XXXX", "XX.XX", "X.XXXX"],
            ["协同增益", "+X.XX", "+X.XXXX", "+X.XX", "+X.XXXX"],
        ],
        col_widths=[3.5, 3.0, 2.8, 3.0, 2.8]
    )

    add_figure_placeholder(doc, "多任务与单任务模型性能对比柱状图")

    add_heading_custom(doc, "4.4  消融实验", level=2)
    add_body_text(doc,
        "为分析各组件对模型性能的贡献，设计了以下消融实验："
    )

    add_table_caption(doc, "消融实验结果")
    add_table_with_style(doc,
        ["实验设置", "疾病分类 Acc (%)", "严重程度 Acc (%)"],
        [
            ["完整模型（MobileNetV2 + 交叉注意力 + 多任务）", "XX.XX", "XX.XX"],
            ["去除交叉注意力（直接使用共享特征）", "XX.XX", "XX.XX"],
            ["去除多任务（单任务模型）", "XX.XX", "XX.XX"],
            ["损失权重 λ₂=0.5", "XX.XX", "XX.XX"],
            ["损失权重 λ₂=1.0", "XX.XX", "XX.XX"],
        ],
        col_widths=[7.0, 3.5, 3.5]
    )

    add_heading_custom(doc, "4.5  可视化分析", level=2)
    add_body_text(doc,
        "本节从训练曲线、Grad-CAM 热力图和混淆矩阵三个维度对模型进行可视化分析。"
    )
    add_figure_placeholder(doc, "训练曲线（损失/准确率/F1 随轮次变化）")
    add_figure_placeholder(doc, "Grad-CAM 热力图对比（正确分类 vs 错误分类案例）")
    add_figure_placeholder(doc, "疾病分类混淆矩阵")

    add_heading_custom(doc, "4.6  少样本学习实验", level=2)
    add_body_text(doc,
        "为评估模型在标注数据有限场景下的泛化能力，本文还进行了少样本学习实验。"
        "在每个类别仅使用 5/10/20 张标注样本的条件下进行训练，测试结果如下。"
    )

    add_table_caption(doc, "少样本学习实验结果")
    add_table_with_style(doc,
        ["样本数/类", "疾病分类 Acc (%)", "疾病分类 F1", "严重程度 Acc (%)"],
        [
            ["5-shot", "XX.XX", "X.XXXX", "XX.XX"],
            ["10-shot", "XX.XX", "X.XXXX", "XX.XX"],
            ["20-shot", "XX.XX", "X.XXXX", "XX.XX"],
            ["Full-shot", "XX.XX", "X.XXXX", "XX.XX"],
        ],
        col_widths=[3.0, 4.0, 4.0, 4.0]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 第 5 章  讨论
    # ══════════════════════════════════════════════════════════════
    add_heading_custom(doc, "第 5 章  讨论", level=1)
    add_body_text(doc,
        "实验结果表明，本文提出的多任务学习框架在农作物病害诊断任务上优于单任务基线模型。"
        "从协同效应角度分析，病害分类和严重程度分级存在互补关系：病害分类需要识别病斑的形态特征，"
        "而严重程度分级需要评估病灶的覆盖范围和扩散程度，这两个任务共享底层视觉特征，"
        "同时各有侧重，适合多任务联合学习。"
    )
    add_body_text(doc,
        "交叉注意力机制的消融实验验证了其有效性。通过为不同任务生成任务特定的注意力权重，"
        "模型能够从共享特征中自适应地提取与各任务相关的信息，避免了硬参数共享带来的特征耦合问题。"
    )
    add_body_text(doc,
        "本文方法也存在一定的局限性。首先，受限于数据集规模，模型对罕见病害的诊断能力有待提升。"
        "其次，当前的严重程度分级仅有三个粗粒度等级，精细化的严重程度评估需要更细粒度的标注数据。"
        "最后，本文的评估仅在实验室条件下进行，真实田间场景下的泛化性能需要进一步验证。"
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 第 6 章  结论与展望
    # ══════════════════════════════════════════════════════════════
    add_heading_custom(doc, "第 6 章  结论与展望", level=1)

    add_heading_custom(doc, "6.1  本文工作总结", level=2)
    add_body_text(doc,
        "本文针对农作物病害智能诊断问题，提出了一个基于多任务深度学习的联合诊断框架。"
        "主要贡献包括：（1）设计了基于 MobileNetV2 和交叉注意力机制的多任务网络结构，"
        "实现了病害分类、严重程度分级和风险评分的联合输出；"
        "（2）通过对比实验验证了多任务学习相较单任务学习的协同优势；"
        "（3）集成了 Grad-CAM 可解释性分析和结构化诊断报告，提升了系统的可用性和透明性；"
        "（4）构建了完整的 Web 诊断系统，具备良好的实际应用价值。"
    )

    add_heading_custom(doc, "6.2  未来工作展望", level=2)
    future_work = [
        "引入 Vision Transformer 等先进架构替代 MobileNetV2 骨干网络，进一步提升分类精度。",
        "研究知识蒸馏技术，将大模型的知识迁移至超轻量级模型，支持移动端实时推理。",
        "融合气象、土壤等多模态数据，实现更全面的作物健康评估。",
        "构建更大规模、更细粒度的农作物病害数据集，覆盖更多作物种类和病害类型。",
        "在实际农田场景中进行部署验证，评估系统在复杂环境下的泛化性能。",
    ]
    for i, fw in enumerate(future_work, 1):
        add_body_text(doc, f"（{i}）{fw}")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 参考文献
    # ══════════════════════════════════════════════════════════════
    add_heading_custom(doc, "参考文献", level=1)

    references = [
        "[1] Mohanty S P, Hughes D P, Salathé M. Using deep learning for image-based plant disease detection[J]. Frontiers in Plant Science, 2016, 7: 1419.",
        "[2] Too E C, Yujian L, Njuki S, et al. A comparative study of fine-tuning deep learning models for plant disease identification[J]. Computers and Electronics in Agriculture, 2019, 161: 272-279.",
        "[3] Sandler M, Howard A, Zhu M, et al. MobileNetV2: Inverted residuals and linear bottlenecks[C]. CVPR, 2018.",
        "[4] Caruana R. Multitask learning[J]. Machine Learning, 1997, 28(1): 41-75.",
        "[5] Misra I, Shrivastava A, Gupta A, et al. Cross-stitch networks for multi-task learning[C]. CVPR, 2016.",
        "[6] Liu S, Johns E, Davison A J. End-to-end multi-task learning with attention[C]. CVPR, 2019.",
        "[7] Selvaraju R R, Cogswell M, Das A, et al. Grad-CAM: Visual explanations from deep networks via gradient-based localization[C]. ICCV, 2017.",
        "[8] Howard A G, Zhu M, Chen B, et al. MobileNets: Efficient convolutional neural networks for mobile vision applications[J]. arXiv:1704.04861, 2017.",
        "[9] He K, Zhang X, Ren S, et al. Deep residual learning for image recognition[C]. CVPR, 2016.",
        "[10] Tan M, Le Q V. EfficientNet: Rethinking model scaling for convolutional neural networks[C]. ICML, 2019.",
        "[11] Dosovitskiy A, Beyer L, Kolesnikov A, et al. An image is worth 16x16 words: Transformers for image recognition at scale[C]. ICLR, 2021.",
        "[12] Hughes D P, Salathé M. An open access repository of images on plant health to enable the development of mobile disease diagnostics[J]. arXiv:1511.08060, 2015.",
    ]

    for ref in references:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.first_line_indent = Cm(0)
        # 悬挂缩进（通过设置左缩进和首行负缩进实现）
        para.paragraph_format.left_indent = Cm(0.74)
        para.paragraph_format.first_line_indent = Cm(-0.74)
        run = para.add_run(ref)
        set_run_font(run, "宋体", "Times New Roman", 10.5)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 致谢
    # ══════════════════════════════════════════════════════════════
    add_heading_custom(doc, "致  谢", level=1)
    add_body_text(doc,
        "（在此处表达对指导教师、同学、家人等的感谢。）"
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 附录
    # ══════════════════════════════════════════════════════════════
    add_heading_custom(doc, "附录 A  主要代码结构说明", level=1)
    add_body_text(doc,
        "本文项目代码组织如下："
    )
    code_structure = [
        ("app.py", "Flask Web 应用主程序，提供前端界面和 REST API"),
        ("train_multitask_model.py", "多任务模型训练脚本，包含数据集、网络结构、训练器、协同效应计算"),
        ("train_single_disease.py", "单任务疾病分类模型训练脚本"),
        ("risk_assessment.py", "风险评估与 Grad-CAM 可视化脚本"),
        ("few_shot_classification.py", "少样本学习实验脚本"),
        ("scripts/utils.py", "共享工具函数（数据加载、标注解析、映射等）"),
        ("scripts/disease_catalog.py", "61 类病害的详细信息目录"),
        ("scripts/export_onnx.py", "模型 ONNX 导出脚本"),
    ]
    add_table_with_style(doc,
        ["文件", "功能说明"],
        code_structure,
        col_widths=[5.0, 10.0]
    )

    # ── 保存文件 ──────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"✅ 论文模板已生成：{OUTPUT_PATH}")


if __name__ == "__main__":
    generate_template()
