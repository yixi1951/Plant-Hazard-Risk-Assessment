"""
完整论文 Word 生成管道
1. 用 pandoc 将 Markdown 转换为 DOCX（原生 OMML 公式）
2. 用 python-docx 后处理：标题页、Team# 页眉、页码、页面设置
"""

import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============ 配置 ============
PROJECT_ROOT = Path(__file__).resolve().parent.parent
MD_PATH = PROJECT_ROOT / "docs" / "论文_基于多任务深度学习的农作物病害智能诊断系统.md"
OUTPUT_DIR = PROJECT_ROOT / "docs"
OUTPUT_NAME = "论文_基于多任务深度学习的农作物病害智能诊断系统.docx"
PANDOC_PATH = r"C:\Program Files\Pandoc\pandoc.exe"

TEAM_NUMBER = "SZUCM2026001"

# ============ 辅助函数 ============

def set_cell_shading(cell, color):
    """设置表格单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name_cn="宋体", name_en="Times New Roman", size=12,
                 bold=False, italic=False, color=None):
    """设置 run 的字体属性"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = name_en
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph_centered(doc, text, font_cn="黑体", font_en="Times New Roman",
                           size=26, bold=True, space_before=72, space_after=6):
    """添加居中段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, font_cn, font_en, size, bold)
    return p


def add_header_footer(doc):
    """为所有节添加 Team# 页眉和页码"""

    def setup_section(section):
        # ---- 页眉 ----
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(f"Team#{TEAM_NUMBER}")
        set_run_font(run, "宋体", "Times New Roman", 9, color=(128, 128, 128))

        # ---- 页脚 ----
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # "— X —" 格式页码
        run1 = fp.add_run("— ")
        set_run_font(run1, "宋体", "Times New Roman", 9, color=(128, 128, 128))

        # PAGE 域代码
        run2 = fp.add_run()
        set_run_font(run2, "Times New Roman", "Times New Roman", 9, color=(128, 128, 128))
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run2._element.append(fldChar1)
        run3 = fp.add_run()
        set_run_font(run3, "Times New Roman", "Times New Roman", 9, color=(128, 128, 128))
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run3._element.append(instrText)
        run4 = fp.add_run()
        set_run_font(run4, "Times New Roman", "Times New Roman", 9, color=(128, 128, 128))
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run4._element.append(fldChar2)

        run5 = fp.add_run(" —")
        set_run_font(run5, "宋体", "Times New Roman", 9, color=(128, 128, 128))

    for section in doc.sections:
        setup_section(section)


def set_page_margins(doc):
    """设置页面边距"""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def clean_team_header_from_content(doc):
    """移除 pandoc 误将第一行 Team# 转换的段落"""
    paragraphs_to_remove = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith(f"Team#{TEAM_NUMBER}") and "Page" in text:
            paragraphs_to_remove.append(p)

    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)

    # 也移除紧随 Team# 行之后的空段（如果有）
    return doc


def fix_image_sizes(doc):
    """适当调整图片大小"""
    for p in doc.paragraphs:
        for run in p.runs:
            for drawing in run._element.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                extents = drawing.findall('.//' + '{http://schemas.openxmlformats.org/drawingml/2006/main}ext')
                for ext in extents:
                    # 限制最大宽度为 14 cm
                    cx = ext.get('cx')
                    if cx:
                        cx_val = int(cx)
                        max_cx = 5333850  # 约 14cm in EMU
                        if cx_val > max_cx:
                            ext.set('cx', str(max_cx))


def enhance_first_para_style(doc):
    """给 pandoc 转换后的第一个段落（即标题）应用大号加粗居中"""
    first_para = None
    for p in doc.paragraphs:
        if p.text.strip():
            first_para = p
            break

    if first_para:
        text = first_para.text.strip()
        # 清空原有内容
        for run in first_para.runs:
            run.text = ""
        first_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        first_para.paragraph_format.space_before = Pt(72)
        first_para.paragraph_format.space_after = Pt(24)
        run = first_para.add_run(text)
        set_run_font(run, "黑体", "Times New Roman", 26, bold=True)
        # 换行
        first_para.add_run("\n").font.size = Pt(12)
        run2 = first_para.add_run(f"Team#{TEAM_NUMBER}")
        set_run_font(run2, "Times New Roman", "Times New Roman", 16, bold=True)

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(12)
        p_sub.paragraph_format.space_after = Pt(36)
        run_sub = p_sub.add_run("机器学习课程大作业")
        set_run_font(run_sub, "宋体", "Times New Roman", 14)


def build_document():
    """主构建流程"""
    print("=" * 60)
    print("论文 Word 文档生成管道")
    print("=" * 60)

    # 1. 准备清理后的临时 Markdown
    print("\n[1/4] 准备 Markdown 源...")
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md_content = f.read()

    # 移除第一行 Team# header 和紧随的 ---
    lines = md_content.split("\n")
    clean_lines = []
    skip_header = True
    skip_dash = True
    for line in lines:
        if skip_header and line.strip().startswith("Team#"):
            skip_header = False
            continue
        if skip_dash and line.strip() == "---":
            skip_dash = False
            continue
        clean_lines.append(line)
    clean_md = "\n".join(clean_lines)

    tmp_md = PROJECT_ROOT / "docs" / "__temp_paper.md"
    with open(tmp_md, "w", encoding="utf-8") as f:
        f.write(clean_md)

    # 2. Pandoc 转换
    print("\n[2/4] Pandoc 转换 Markdown → DOCX（含 OMML 公式）...")
    tmp_docx = PROJECT_ROOT / "docs" / "__temp_paper.docx"

    pandoc_cmd = [
        str(PANDOC_PATH),
        str(tmp_md),
        "-o", str(tmp_docx),
        "--from", "markdown+tex_math_dollars",
        "--to", "docx",
        "--resource-path", f"{PROJECT_ROOT / 'docs'};{PROJECT_ROOT / 'docs' / 'img'}",
        "--metadata", "title=基于多任务深度学习的农作物病害智能诊断系统",
        "--wrap=none",
    ]

    result = subprocess.run(pandoc_cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"Pandoc 错误: {result.stderr}")
        raise RuntimeError("Pandoc 转换失败")

    # 3. 后处理
    print("\n[3/4] 后处理 — 标题页、页眉页码、格式调整...")
    doc = Document(str(tmp_docx))

    # 清理 Team# header 段落
    clean_team_header_from_content(doc)

    # 插入标题页
    # 在文档最前面插入标题页
    title_p = doc.paragraphs[0]._element
    parent = title_p.getparent()

    # 创建标题段落
    new_p1 = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr><w:jc w:val="center"/><w:spacing w:before="1440" w:after="240"/></w:pPr>'
        f'  <w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:eastAsia="黑体"/>'
        f'    <w:b/><w:sz w:val="52"/></w:rPr>'
        f'  <w:t xml:space="preserve">基于多任务深度学习的\n农作物病害智能诊断系统</w:t></w:r>'
        f'</w:p>'
    )
    parent.insert(0, new_p1)

    new_p2 = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr><w:jc w:val="center"/><w:spacing w:before="480" w:after="120"/></w:pPr>'
        f'  <w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:eastAsia="Times New Roman"/>'
        f'    <w:b/><w:sz w:val="32"/></w:rPr>'
        f'  <w:t>Team#{TEAM_NUMBER}</w:t></w:r>'
        f'</w:p>'
    )
    # 在标题段落之后插入
    parent.insert(1, new_p2)

    new_p3 = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr><w:jc w:val="center"/><w:spacing w:before="240" w:after="720"/></w:pPr>'
        f'  <w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:eastAsia="宋体"/>'
        f'    <w:sz w:val="28"/></w:rPr>'
        f'  <w:t>机器学习课程大作业</w:t></w:r>'
        f'</w:p>'
    )
    parent.insert(2, new_p3)

    # 分页
    page_break = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:r><w:br w:type="page"/></w:r>'
        f'</w:p>'
    )
    parent.insert(3, page_break)

    # 设置页面边距
    set_page_margins(doc)

    # 添加页眉页码
    add_header_footer(doc)

    # 调整图片大小
    fix_image_sizes(doc)

    # 4. 保存
    print("\n[4/4] 保存最终文档...")
    output_path = OUTPUT_DIR / OUTPUT_NAME
    doc.save(str(output_path))

    # 清理临时文件
    if tmp_md.exists():
        tmp_md.unlink()
    if tmp_docx.exists():
        tmp_docx.unlink()

    # 统计
    print(f"\n{'=' * 60}")
    print(f"✓ 成功生成: {output_path}")
    file_size = output_path.stat().st_size
    print(f"  文件大小: {file_size / 1024:.1f} KB")

    # 验证
    doc2 = Document(str(output_path))
    para_count = len(doc2.paragraphs)
    table_count = len(doc2.tables)

    # 统计 OMML 公式
    omath_count = 0
    with zipfile.ZipFile(str(output_path)) as z:
        if 'word/document.xml' in z.namelist():
            xml_content = z.read('word/document.xml').decode('utf-8')
            omath_count += xml_content.count('m:oMath')
            omath_count += xml_content.count('m:oMathPara')

    img_count = 0
    for p in doc2.paragraphs:
        for run in p.runs:
            if run._element.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                img_count += 1

    print(f"  段落: {para_count} | 表格: {table_count} | 公式: {omath_count} | 图片: {img_count}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    import zipfile
    build_document()
