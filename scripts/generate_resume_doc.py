# -*- coding: utf-8 -*-
"""生成简历项目简介与面试问答 Word 文档。"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from pathlib import Path


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=16 if level == 1 else 14, bold=True)
    return h


def add_para(doc, text, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(4)
    return p


def add_qa(doc, question, answer):
    pq = doc.add_paragraph()
    rq = pq.add_run(f"Q：{question}")
    set_run_font(rq, size=11, bold=True, color=(27, 67, 50))
    pq.paragraph_format.space_before = Pt(10)
    pq.paragraph_format.space_after = Pt(4)

    pa = doc.add_paragraph()
    ra = pa.add_run(f"A：{answer}")
    set_run_font(ra, size=10.5)
    pa.paragraph_format.left_indent = Cm(0.5)
    pa.paragraph_format.line_spacing = 1.4
    pa.paragraph_format.space_after = Pt(8)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 封面标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("智农 · 农作物病虫害 AI 智能识别预警系统")
    set_run_font(tr, size=20, bold=True, color=(27, 67, 50))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("简历项目简介 · 面试问答背诵手册")
    set_run_font(sr, size=12, color=(92, 83, 71))

    doc.add_paragraph()

    # ========== 一、简历可直接粘贴 ==========
    add_heading(doc, "一、简历可直接粘贴的内容", 1)

    add_heading(doc, "1.1 项目名称（推荐）", 2)
    add_para(doc, "智农 — 农作物病虫害 AI 智能识别与预警系统（个人全栈项目）")

    add_heading(doc, "1.2 一句话版本（项目列表旁）", 2)
    add_para(
        doc,
        "基于 MobileNetV2 多任务深度学习，实现叶片病害 61 类识别、严重程度分级与风险预警，"
        "并搭建 Flask Web 控制台完成推理部署、批量诊断与 JSON/PDF 报告导出。",
    )

    add_heading(doc, "1.3 简历 bullet 版（推荐 3～4 条）", 2)
    bullets = [
        "独立设计并实现多任务 CNN（MobileNetV2 + CrossTaskAttention），在 PlantVillage 61 类验证集上 Top-1 准确率约 74.7%，严重程度分类准确率约 85%+。",
        "搭建 Flask 推理服务与可视化控制台，支持拖拽/URL/批量 ZIP 上传、ECharts 风险趋势图、Grad-CAM 热力图与 MC Dropout 不确定性分析。",
        "基于病害知识库与规则引擎生成可执行防治方案（分步路线图、复查时间线、成本估算），并自动导出 JSON/PDF 诊断报告。",
        "完成数据准备（PlantVillage/Kaggle 适配）、模型训练脚本、演示数据注入与 REST API，可脱离真实模型向客户演示完整业务闭环。",
    ]
    for b in bullets:
        add_bullet(doc, b)

    add_heading(doc, "1.4 技术栈关键词（Skills 区可写）", 2)
    add_para(doc, "Python · PyTorch · MobileNetV2 · 多任务学习 · Flask · ECharts · PIL/OpenCV · scikit-learn · PlantVillage")

    add_heading(doc, "1.5 谨慎表述（避免面试官追问穿帮）", 2)
    cautions = [
        "❌ 不要写「双模型融合」「准确率 92%」「第二模型生成防治方案」—— 防治方案是知识库 + 规则引擎，不是独立神经网络。",
        "❌ 不要写「已商用上线」「服务 XX 万用户」—— 这是个人学习与演示项目。",
        "❌ 不要写「实时视频流识别」—— 当前是单张/批量图片推理。",
        "✅ 应写「验证集 Top-1 ~74.7%」「个人全栈」「可演示完整流程」「MC Dropout 不确定性量化」。",
        "✅ 被问到不足时主动说：61 类细粒度分类仍有提升空间，后续可加更强 backbone、数据增强、模型蒸馏或 ONNX 边缘部署。",
    ]
    for c in cautions:
        add_bullet(doc, c)

    # ========== 二、项目深度介绍 ==========
    add_heading(doc, "二、项目深度介绍（30 秒 / 1 分钟 / 3 分钟）", 1)

    add_heading(doc, "2.1 30 秒电梯演讲", 2)
    add_para(
        doc,
        "这是一个面向农业场景的个人全栈 AI 项目。用户上传作物叶片照片，系统用多任务深度学习模型"
        "同时识别病害种类和严重程度，给出风险评分；再结合病害知识库输出防治建议和诊断报告。"
        "我负责从数据处理、模型训练到 Flask 部署和前端控制台全部环节，验证集 61 类 Top-1 约 74.7%。",
    )

    add_heading(doc, "2.2 1 分钟版本", 2)
    add_para(
        doc,
        "背景：农户和农技人员靠经验判断病害，效率低、误判成本高。"
        "方案：我用 PlantVillage 公开数据集训练 MobileNetV2 多任务网络，共享骨干 + CrossTaskAttention，"
        "一个模型同时输出 61 类病害分类和 3 级严重程度。"
        "工程：Flask 提供上传识别、批量预测、REST API；前端有趋势图、结果可视化、报告导出；"
        "防治方案由规则引擎根据识别结果匹配知识库生成。"
        "结果：Top-1 约 74.7%，宏 F1 约 0.70，严重程度约 85%+；并做了 MC Dropout 和 Grad-CAM 增强可信度。",
    )

    add_heading(doc, "2.3 3 分钟版本（按 STAR 结构）", 2)
    add_para(doc, "【Situation 背景】", bold=True)
    add_para(
        doc,
        "农业病害识别依赖人工经验，PlantVillage 等公开数据已证明 CNN 可行，但多数 Demo 只做分类，"
        "缺少严重程度、风险预警和可落地防治建议的完整闭环。",
    )
    add_para(doc, "【Task 任务】", bold=True)
    add_para(
        doc,
        "作为个人项目，我需要完成：① 多任务模型训练与评估；② Web 推理服务；③ 可演示的业务流程（含报告与防治方案）。",
    )
    add_para(doc, "【Action 行动】", bold=True)
    add_bullet(doc, "数据：PlantVillage 61 类 + 模拟数据快速验证；脚本支持 Kaggle 目录适配。")
    add_bullet(doc, "模型：MobileNetV2 骨干 + CrossTaskAttention + 双分类头；128×128 输入，兼顾速度与精度。")
    add_bullet(doc, "推理：predict_image / predict_with_uncertainty（MC Dropout 16 次采样）。")
    add_bullet(doc, "后端：Flask 路由、批量 ZIP、临时文件清理、URL 抓图安全校验。")
    add_bullet(doc, "业务：treatment_engine 规则引擎匹配 DISEASE_DETAILS 知识库，生成路线图与时间线。")
    add_bullet(doc, "演示：demo_data 注入 7 条历史记录 + 4 个样例，无模型权重也能展示给客户。")
    add_para(doc, "【Result 结果】", bold=True)
    add_para(
        doc,
        "验证集 Top-1 74.73%（Epoch 30），宏 F1 约 0.70；Web 端可完成上传→识别→方案→报告全流程；"
        "项目代码结构清晰，适合作为 CV + 工程化综合展示。",
    )

    # ========== 三、架构与技术细节 ==========
    add_heading(doc, "三、架构与技术细节速记", 1)

    add_heading(doc, "3.1 模型结构（口述用）", 2)
    add_para(
        doc,
        "输入 128×128 RGB → MobileNetV2 特征提取 → AdaptiveAvgPool → CrossTaskAttention "
        "→ 病害头（N 类 Softmax）+ 严重程度头（3 类 Softmax）→ 综合风险评分与置信度。",
    )

    add_heading(doc, "3.2 CrossTaskAttention 怎么说", 2)
    add_para(
        doc,
        "共享 backbone 输出 1280 维特征后，分别做 disease query 和 severity query，"
        "与 value 投影做点积注意力，得到两个任务增强的特征再进各自分类头。"
        "目的是让「是什么病」和「有多严重」互相利用信息，比完全独立两个模型参数更少、推理更快。",
    )

    add_heading(doc, "3.3 防治方案怎么说（重要）", 2)
    add_para(
        doc,
        "防治方案不是第二个深度学习模型，而是 treatment_engine：根据模型输出的病害 ID/名称、"
        "严重程度、置信度，查 DISEASE_DETAILS 知识库，按规则调整 urgency、步骤、成本与复查时间。"
        "优势是可解释、可编辑、农技人员可维护；不足是覆盖范围受知识库限制。",
    )

    add_heading(doc, "3.4 核心 API", 2)
    apis = [
        "POST / — 单张图片识别（action=predict）",
        "POST /batch_predict — 多图或 ZIP 批量预测",
        "GET /api/diseases — 病害列表（含防治信息）",
        "GET /api/dashboard_stats — 仪表盘统计与趋势图数据",
        "GET /reports/<fname> — 下载 JSON/PDF 报告",
    ]
    for a in apis:
        add_bullet(doc, a)

    # ========== 四、面试问答 ==========
    add_heading(doc, "四、高频面试问答（建议背诵）", 1)

    qa_list = [
        (
            "用 1 分钟介绍一下这个项目。",
            "参考第二节 1 分钟版本，先说背景→方案→你的职责→量化结果，控制在 60～90 秒。",
        ),
        (
            "这是你一个人做的吗？团队分工？",
            "是个人全栈项目，从数据处理、模型训练、后端 API、前端界面到演示脚本都是我独立完成的。"
            "简历里如果写了团队，应改为个人项目，避免被追问分工细节。",
        ),
        (
            "为什么用多任务学习，而不是两个独立模型？",
            "病害种类和严重程度在视觉特征上有相关性，共享 backbone 减少参数量和推理时间，"
            "CrossTaskAttention 让两个任务特征互相增强。实验上 severity 头也能到 85%+，"
            "说明共享表示是有效的。若数据极不均衡，也可拆成独立模型对比 ablation。",
        ),
        (
            "CrossTaskAttention 具体怎么实现？",
            "对 pooled 特征分别线性投影为 disease query、severity query 和 value，"
            "用 query 与 value 逐元素相乘后 softmax 得到注意力权重，加权后送入各自 MLP 分类头。"
            "不是 Transformer 多头，而是轻量级 task-specific attention。",
        ),
        (
            "74.7% 准确率算高吗？怎么评估的？",
            "在 PlantVillage 61 类细粒度分类上，这是验证集 Top-1，Epoch 30 最佳约 74.73%，宏 F1 约 0.70。"
            "61 类难度较高，工业界通常还会结合业务场景看召回/F1 和人工复核流程。"
            "我认为仍有提升空间：更强 backbone（EfficientNet）、CutMix、类别重采样、集成学习等。",
        ),
        (
            "数据从哪来？怎么预处理的？",
            "主要用 PlantVillage 公开数据集，也写了 mock 数据脚本做流程验证。"
            "预处理：Resize 128×128、ImageNet 均值方差归一化；训练时有数据增强（翻转、颜色抖动等，视训练脚本配置）。"
            "标签从文件夹名或 TXT 标注解析，severity 可由病害名规则映射或单独标注。",
        ),
        (
            "MC Dropout 是什么？你为什么用？",
            "推理时保持 Dropout 开启，对同一张图多次前向采样，得到预测均值和方差。"
            "方差大说明模型不确定，我在 UI 上提示人工复核，避免低置信度误判直接指导用药。",
        ),
        (
            "Grad-CAM 做了什么？",
            "对 severity 或 disease 分支做梯度加权类激活映射，热力图叠加在原图上，"
            "展示模型关注区域，帮助用户理解「为什么判这个病」，增强可解释性。",
        ),
        (
            "防治方案是怎么生成的？是 AI 吗？",
            "是规则引擎 + 结构化知识库，不是生成式大模型也不是第二个 CNN。"
            "流程：模型输出病害名/ID 和 severity → 查 DISEASE_DETAILS → treatment_engine 组装步骤、"
            "urgency、成本、季节提示；若 MC Dropout std 高则附加 uncertainty_note。诚实说这点反而加分。",
        ),
        (
            "Flask 服务怎么设计的？并发怎么处理？",
            "Flask 同步路由，启动时 preload 模型到内存；predict 走 PyTorch no_grad。"
            "开发用 app.run，生产可换 Gunicorn 多 worker。"
            "上传有限制 MAX_CONTENT_LENGTH，tmp_uploads 定时清理；批量预测用 FormData 多文件或 ZIP 解压逐张推理。",
        ),
        (
            "项目最难的点是什么？怎么解决的？",
            "难点 1：多任务标签对齐（61 类 disease + severity），用统一 Dataset 类和 metadata JSON 管理 class_names。"
            "难点 2：演示环境常无 GPU/权重，用 demo_data 注入样例报告，保证面试/客户演示不断链。"
            "难点 3：前端趋势图 tooltip 等交互 bug，关闭 ECharts 默认 tooltip 改自定义层。",
        ),
        (
            "如果上线到真实农田，还要做什么？",
            "① 收集本地作物数据微调；② 模型量化/ONNX/TensorRT 边缘部署；③ 异步任务队列处理批量；"
            "④ 用户鉴权与日志审计；⑤ 低置信度强制人工复核；⑥ 与农技专家共建并审核防治知识库。",
        ),
        (
            "为什么选 MobileNetV2？",
            "轻量、速度快，适合 CPU/边缘设备 demo；128 输入进一步降低延迟。"
            "若追求精度可换 EfficientNet-B0 或 ConvNeXt-Tiny，但要重新权衡推理时延。",
        ),
        (
            "训练和推理环境？",
            "训练：PyTorch + CUDA（如有）；推理：自动检测 cuda/cpu。"
            "依赖见 requirements.txt：torch、torchvision、Flask、PIL、opencv、sklearn 等。",
        ),
        (
            "批量预测怎么实现？",
            "前端 FormData 传多文件或 ZIP，后端 zipfile 解压，循环调用 predict_image，"
            "汇总 JSON 返回每张图的 disease_name、severity、risk_score、treatment_summary。",
        ),
        (
            "你如何证明项目是你做的？",
            "能讲清 CrossTaskAttention 结构、训练脚本参数、Flask 路由、"
            "treatment_engine 规则逻辑；本地可 live demo 上传识别；"
            "能解释验证集指标来源（训练日志 / plot_training_curves 中 Epoch 30, 74.73%）。",
        ),
        (
            "这个项目有什么不足？",
            "① 61 类准确率仍有提升空间；② 防治方案依赖静态知识库，未覆盖所有地域用药规范；"
            "③ 未做移动端适配和正式生产监控；④ 类别不平衡时少数类召回可能偏低。主动说不足显成熟。",
        ),
        (
            "和 YOLO/目标检测方案比呢？",
            "当前是分类范式，假设叶片/病斑已在画面中；若需田间复杂背景定位，"
            "可升级为检测+分类两阶段或弱监督定位。分类方案实现快、适合叶片特写场景。",
        ),
        (
            "面试官让现场演示怎么办？",
            "提前 py app.py 启动，打开 localhost:7860；优先点「客户演示样例」展示完整结果；"
            "再上传一张清晰叶片图；展示趋势图、防治方案、下载报告。"
            "若无 GPU，强调 CPU 也可推理，演示数据不依赖权重。",
        ),
    ]

    for q, a in qa_list:
        add_qa(doc, q, a)

    # ========== 五、追问应对 ==========
    add_heading(doc, "五、刁钻追问 & 标准应答", 1)

    add_qa(
        doc,
        "92% 准确率是你说的吗？",
        "不是。本项目验证集 Top-1 约 74.7%，我不会夸大。"
        "若简历旧版写过 92%，应改为真实数字或注明「严重程度子任务」以免混淆。",
    )
    add_qa(
        doc,
        "第二个模型在哪里？",
        "没有第二个预测模型。防治方案是 treatment_engine 规则系统；"
        "只有一个 MultiTaskNetwork 负责 disease + severity 分类。",
    )
    add_qa(
        doc,
        "Attention 是不是噱头？",
        "可以坦诚：这是 task-level 轻量 attention，不是大模型 Attention。"
        "价值在于多任务特征解耦与增强，若重做会加 ablation（有/无 attention 对比）用实验说话。",
    )
    add_qa(
        doc,
        "为什么不用大模型 / GPT？",
        "病害识别是视觉分类任务，小模型在 PlantVillage 上更成熟、可本地部署、成本低。"
        "GPT 可用于报告文案润色，但核心识别必须靠 CV 模型保证可控与可复现。",
    )

    # ========== 六、背诵 checklist ==========
    add_heading(doc, "六、面试前 10 分钟 Checklist", 1)
    checklist = [
        "能背出：MobileNetV2 + CrossTaskAttention + 双头，128×128 输入",
        "能背出：Top-1 ~74.7%（61 类），severity ~85%+，宏 F1 ~0.70",
        "能背出：Flask + ECharts + MC Dropout + Grad-CAM",
        "能背出：防治方案 = 知识库 + 规则引擎（非第二模型）",
        "能背出：个人全栈，PlantVillage 公开数据",
        "本地能打开 http://127.0.0.1:7860 演示",
        "准备 1 张自己上传识别的截图 / 报告 JSON",
    ]
    for i, item in enumerate(checklist, 1):
        add_bullet(doc, f"{i}. {item}")

    add_para(doc, "")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("— 祝面试顺利 —")
    set_run_font(fr, size=10, color=(92, 83, 71))

    return doc


def main():
    out = Path(__file__).resolve().parents[1] / "docs" / "智农项目_简历与面试问答.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(str(out))
    print(f"已生成: {out}")


if __name__ == "__main__":
    main()
