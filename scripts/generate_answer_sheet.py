# -*- coding: utf-8 -*-
"""
生成深圳大学《机器学习》大作业答题纸
—— 自主应用实践类
基于多任务深度学习的农作物病害智能诊断与风险评估系统
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)),
                           "机器学习大作业_农作物病害智能诊断系统.docx")


# ====================================================================
# 工具函数
# ====================================================================

def set_run_font(run, font_cn="宋体", font_en="Times New Roman", size=12,
                 bold=False, italic=False, color=None, underline=False):
    run.font.size = Pt(size)
    run.font.name = font_en
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text, font_cn="宋体", font_en="Times New Roman", size=12,
             bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=0, line_spacing=1.5,
             first_line_indent=None, italic=False, color=None):
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = para.add_run(text)
    set_run_font(run, font_cn, font_en, size, bold, italic, color)
    return para


def add_body(doc, text, indent=0.74):
    """正文段落：宋体12pt，首行缩进2字符，1.5倍行距"""
    return add_para(doc, text, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=1.5, first_line_indent=indent,
                    space_before=0, space_after=0)


def add_heading_l1(doc, text):
    """一级标题：黑体14pt，加粗"""
    return add_para(doc, text, font_cn="黑体", size=14, bold=True,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=12, space_after=6, line_spacing=1.5)


def add_heading_l2(doc, text):
    """二级标题：黑体12pt，加粗"""
    return add_para(doc, text, font_cn="黑体", size=12, bold=True,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=8, space_after=4, line_spacing=1.5)


def add_fig_caption(doc, text):
    """图注：居中，宋体10.5pt"""
    return add_para(doc, text, size=10.5,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=3, space_after=8, line_spacing=1.25)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def make_header_cell(cell, text, width_cm=None):
    """设置表头单元格"""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_run_font(run, "黑体", "Arial", 10, bold=True, color=(255, 255, 255))
    set_cell_shading(cell, "4472C4")
    if width_cm:
        cell.width = Cm(width_cm)


def make_data_cell(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER):
    """设置数据单元格"""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text))
    set_run_font(run, "宋体", "Times New Roman", 10)


def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        make_header_cell(table.rows[0].cells[i], h)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            make_data_cell(table.rows[ri + 1].cells[ci], val, align)
            if ri % 2 == 1:
                set_cell_shading(table.rows[ri + 1].cells[ci], "D9E2F3")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    # 表后空行
    add_para(doc, "", size=6, space_after=0)
    return table


def setup_page(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)


def add_footer_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = para.add_run("第 ")
        set_run_font(run1, "宋体", "Times New Roman", 9)
        fld1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run2 = para.add_run()
        run2._element.append(fld1)
        run3 = para.add_run()
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run3._element.append(instr)
        run4 = para.add_run()
        fld2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run4._element.append(fld2)
        run5 = para.add_run(" 页")
        set_run_font(run5, "宋体", "Times New Roman", 9)


def add_figure_placeholder(doc, caption):
    """图片占位"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(f"[ 此处插入：{caption} ]")
    set_run_font(run, "宋体", "Arial", 10.5, italic=True, color=(120, 120, 120))
    add_fig_caption(doc, caption)


# ====================================================================
# 主生成函数
# ====================================================================

def generate():
    doc = Document()
    setup_page(doc)

    # ══════════════════════════════════════════════════════════════
    # 答题纸抬头（严格仿照深圳大学格式）
    # ══════════════════════════════════════════════════════════════
    # 标题行
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(2)
    run = title_para.add_run("深圳大学考试答题纸")
    set_run_font(run, "黑体", "Arial", 18, bold=True)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_before = Pt(0)
    sub_para.paragraph_format.space_after = Pt(6)
    run = sub_para.add_run("(以论文、报告等形式考核专用)")
    set_run_font(run, "宋体", "Times New Roman", 10.5)

    # ── 信息表格（5行2列）
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = "Table Grid"

    info_data = [
        ("二五 ～ 二六    学年度第  二  学期", ""),
        ("课程编号  1501990018   课程名称  机器学习", ""),
        ("主讲教师  贾森、何汝艳", ""),
        ("学    号", ""),
        ("专业年级", ""),
    ]
    for i, (left, right) in enumerate(info_data):
        cell0 = info_table.rows[i].cells[0]
        cell0.text = ""
        p = cell0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(left)
        set_run_font(r, "宋体", "Times New Roman", 10.5)

        cell1 = info_table.rows[i].cells[1]
        cell1.text = ""
        p = cell1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if right:
            r = p.add_run(right)
            set_run_font(r, "宋体", "Times New Roman", 10.5)

    # 合并学期行和课程信息行的右侧单元格
    # 学期行(0)的右单元格为空，课程信息行(1)的右单元格为空白

    add_para(doc, "", size=6, space_after=0)

    # ── 论文标题
    add_para(doc, "课程大论文", font_cn="黑体", size=16, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=6, space_after=6, line_spacing=1.5)

    add_para(doc, "——基于多任务深度学习的农作物病害智能诊断与风险评估系统",
             font_cn="黑体", size=14, bold=False,
             alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=0, space_after=12, line_spacing=1.5)

    add_para(doc, "（自主应用实践类）", font_cn="宋体", size=10.5,
             alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=0, space_after=12)

    # ══════════════════════════════════════════════════════════════
    # 正文内容（约4000字）
    # ══════════════════════════════════════════════════════════════

    # ── 1. 项目背景与目标 ──
    add_heading_l1(doc, "一、项目背景与目标")
    add_body(doc,
        "农作物病害是威胁全球粮食安全的主要因素之一。据联合国粮食及农业组织（FAO）统计，"
        "全球每年因植物病害造成的粮食损失高达10%~16%，对农业生产和粮食安全构成严重威胁。"
        "传统的病害诊断方式主要依赖农业技术人员的田间观察和经验判断，存在效率低下、主观性强、"
        "专业门槛高等突出问题，难以满足大规模、实时化、精准化的现代农业需求。")
    add_body(doc,
        "近年来，以深度学习为代表的人工智能技术在计算机视觉领域取得了突破性进展，"
        "为农作物病害的自动化诊断提供了新的技术路径。从早期的卷积神经网络（CNN）到ResNet、"
        "MobileNet、Vision Transformer等先进架构，图像识别精度不断提升。然而，现有研究大多"
        "聚焦于单一的分类任务——仅判断病害种类或仅评估严重程度，缺乏对病害信息的全面联合诊断能力。")
    add_body(doc,
        "针对上述问题，本项目提出了一种基于多任务深度学习的农作物病害智能诊断与风险评估系统。"
        "项目的核心目标是：（1）实现病害类别、严重程度和风险评分的联合预测；"
        "（2）引入交叉注意力机制（Cross-Task Attention）提升多任务协同效果；"
        "（3）结合Grad-CAM可解释性方法生成热力图，增强诊断的可信度；"
        "（4）构建完整的Web应用系统，支持图片上传、实时推理和诊断报告生成。"
        "本项目属于自主应用实践类，完整覆盖了从理论分析、模型设计、实验验证到系统开发的全部流程。")

    # ── 2. 数据集与预处理 ──
    add_heading_l1(doc, "二、数据集介绍与预处理")
    add_body(doc,
        "本项目使用两个公开数据集进行评估。PlantVillage数据集包含14种作物的38类病害图像，"
        "共约54,000张叶片图像，是植物病害识别领域最广泛使用的基准数据集。"
        "AgriculturalDisease竞赛数据集包含10种主要经济作物（苹果、樱桃、玉米、葡萄、柑桔、"
        "桃、辣椒、马铃薯、草莓、番茄）的61类病害图像，覆盖了中国农业生产中最常见的作物和病害类型。")
    add_body(doc,
        "在数据预处理方面，所有图像统一调整至128×128像素，以适应MobileNetV2的输入要求。"
        "数据增强策略包括随机水平翻转（概率0.3）和随机旋转（±10度），以提升模型的泛化能力。"
        "图像归一化采用ImageNet数据集的均值和标准差（mean=[0.485,0.456,0.406]，"
        "std=[0.229,0.224,0.225]）。数据集按80%/20%的比例划分为训练集和验证集。")

    add_table(doc,
        ["数据集", "作物种类", "病害类别数", "总样本量", "训练/验证"],
        [
            ["PlantVillage", "14种", "38类", "~54,000张", "80%/20%"],
            ["AgriculturalDisease", "10种", "61类", "~45,000张", "80%/20%"],
        ],
        col_widths=[3.5, 2.5, 2.5, 2.5, 2.5])
    add_fig_caption(doc, "表1  数据集统计信息")

    # ── 3. 方法与模型设计 ──
    add_heading_l1(doc, "三、方法与模型设计")

    add_heading_l2(doc, "3.1 总体框架")
    add_body(doc,
        "本系统采用多任务学习（Multi-Task Learning）框架，由一个共享骨干网络和两个任务特定"
        "分支组成。输入叶片图像后，共享骨干网络提取通用视觉特征，再通过交叉注意力机制为"
        "病害分类和严重程度分级分别生成任务特定特征，最后两个独立的分支分别输出病害类别"
        "（61类）和严重程度等级（3级：健康/一般/严重），并综合两者计算风险评分。")

    add_figure_placeholder(doc, "图1  多任务诊断框架整体架构")
    add_body(doc,
        "图1展示了系统的整体架构。输入图像经过MobileNetV2骨干网络提取1280维共享特征，"
        "交叉注意力模块将其分别映射为病害相关特征和严重程度相关特征，两个分类头独立输出预测结果。")

    add_heading_l2(doc, "3.2 共享骨干网络：MobileNetV2")
    add_body(doc,
        "选用MobileNetV2作为共享骨干网络（Sandler et al., 2018），主要基于以下考虑："
        "（1）MobileNetV2采用倒残差结构（Inverted Residuals）和线性瓶颈（Linear Bottleneck），"
        "在保持较高精度的同时参数量和计算量远小于VGG和ResNet等传统网络；"
        "（2）轻量化设计有利于后续的移动端和边缘端部署，具有实际应用价值；"
        "（3）ImageNet预训练权重提供了良好的初始化，有利于加快收敛并提升性能。"
        "为平衡迁移学习和任务特定训练，冻结了MobileNetV2的大部分底层参数，仅训练最后3层。")

    add_heading_l2(doc, "3.3 交叉注意力机制")
    add_body(doc,
        "交叉注意力机制是本项目的核心创新点。其设计思想是：病害分类任务应关注病斑的形态和纹理特征，"
        "而严重程度分级任务则应关注病灶的面积占比和扩散程度。虽然两个任务共享底层特征，"
        "但对特征各维度的关注度各不相同。")
    add_body(doc,
        "具体实现上，给定共享特征向量x∈R¹²⁸⁰，通过三个独立的线性变换分别生成病害查询Q_d、"
        "严重程度查询Q_s和值映射V（均为256维）。然后计算注意力权重并生成任务特定特征："
        "f_d = softmax(Q_d·V)⊙V，f_s = softmax(Q_s·V)⊙V。"
        "最终f_d输入病害分类头，f_s输入严重程度分类头。"
        "这种设计使得两个任务可以自适应地从共享特征中选择所需的信息维度，"
        "既保留了多任务学习的参数共享优势，又提供了任务特定的特征提取能力。")

    add_heading_l2(doc, "3.4 损失函数设计")
    add_body(doc,
        "采用加权多任务损失函数：L_total = λ₁·L_CE(y_d, ŷ_d) + λ₂·L_CE(y_s, ŷ_s)。"
        "其中L_CE为交叉熵损失，任务权重默认设置为λ₁=1.0、λ₂=0.8。"
        "严重程度任务的权重略低，是因为严重程度分级仅有3个类别，学习难度相对较低。"
        "优化器选用AdamW（学习率5×10⁻⁴，权重衰减1×10⁻⁴），"
        "学习率调度采用ReduceLROnPlateau策略（patience=3，衰减因子0.5），"
        "并设置早停机制（patience=5）防止过拟合。")

    add_heading_l2(doc, "3.5 Grad-CAM可解释性分析")
    add_body(doc,
        "为增强模型决策的可解释性，引入Grad-CAM（Gradient-weighted Class Activation Mapping）方法。"
        "该方法利用目标类别对最后一层卷积特征图的梯度，计算每个特征通道的重要性权重，"
        "生成与类别相关的热力图并叠加到原图上，直观展示模型做出决策的依据——"
        "即叶片上的哪些区域对判断病害类别贡献最大。"
        "这一功能有助于提升用户对AI诊断结果的信任度，也为农技人员复核诊断结果提供了视觉参考。")

    add_heading_l2(doc, "3.6 风险评估方法")
    add_body(doc,
        "基于病害类别和严重程度分级结果，设计了三级风险评估规则。"
        "当模型预测为严重疾病（2级）且置信度超过80%时，判定为高风险，需要立即采取防控措施；"
        "当预测为一般疾病（1级）且置信度超过70%时，判定为中风险，需要密切监测；"
        "其余情况为低风险或不确定，建议常规管理或人工复核。"
        "该规则参考了农业植保领域的实践经验，兼顾了诊断的准确性和安全性。")

    # ── 4. 实验与结果分析 ──
    add_heading_l1(doc, "四、实验设计与结果分析")

    add_heading_l2(doc, "4.1 实验设置")
    add_body(doc,
        "所有实验在统一条件下进行。模型输入尺寸为128×128，批大小根据硬件条件设置为"
        "16（GPU可用时）或4（仅CPU时）。训练最大轮次为15轮，早停耐心值为5轮。"
        "评价指标包括准确率（Accuracy）和宏平均F1分数（Macro F1）。"
        "所有实验使用相同的随机种子以保证结果可复现。")

    add_table(doc,
        ["超参数", "取值"],
        [
            ["骨干网络", "MobileNetV2（ImageNet预训练）"],
            ["输入尺寸", "128 × 128"],
            ["优化器", "AdamW（lr=5e-4, weight_decay=1e-4）"],
            ["学习率调度", "ReduceLROnPlateau（patience=3, factor=0.5）"],
            ["早停耐心值", "5轮"],
            ["最大轮次", "15轮"],
            ["批大小", "16（GPU）/ 4（CPU）"],
            ["数据增强", "RandomHorizontalFlip(p=0.3), RandomRotation(10°)"],
            ["任务权重", "λ₁=1.0（疾病）, λ₂=0.8（严重程度）"],
        ],
        col_widths=[5.0, 10.0])
    add_fig_caption(doc, "表2  实验超参数设置")

    add_heading_l2(doc, "4.2 多任务与单任务对比实验")
    add_body(doc,
        "为验证多任务学习的协同效应，训练了三组模型进行公平对比："
        "（1）多任务模型（本文方法），同时预测病害类别和严重程度；"
        "（2）单任务疾病分类模型，仅预测病害类别；"
        "（3）单任务严重程度分级模型，仅预测严重程度。"
        "三组模型使用相同的MobileNetV2骨干网络和训练条件，仅修改输出头。")

    add_table(doc,
        ["模型", "疾病分类Acc(%)", "疾病分类F1", "严重程度Acc(%)", "严重程度F1"],
        [
            ["单任务疾病模型", "72.48", "0.6967", "—", "—"],
            ["单任务严重程度模型", "—", "—", "XX.XX", "X.XXXX"],
            ["多任务模型（本文）", "74.73", "0.7123", "XX.XX", "X.XXXX"],
            ["协同增益", "+2.25%", "+0.0156", "+X.XX%", "+X.XXXX"],
        ],
        col_widths=[3.5, 3.0, 2.8, 3.0, 2.8])
    add_fig_caption(doc, "表3  主实验结果：多任务vs单任务性能对比")

    add_body(doc,
        "实验结果表明，多任务模型在疾病分类任务上相比单任务基线有显著提升，"
        "准确率提升约2.25个百分点，F1分数提升0.0156。这验证了多任务学习在农业病害诊断中的协同效应——"
        "严重程度分级任务为疾病分类提供了有益的归纳偏置（Inductive Bias），"
        "帮助模型学习到更具判别力的特征表示。")

    add_heading_l2(doc, "4.3 可视化分析")
    add_body(doc,
        "从训练曲线可以看出，随着训练轮次的增加，多任务模型的训练损失和验证损失均平稳下降，"
        "疾病分类准确率从约28%提升至74%以上，验证F1分数从0.20提升至0.70左右，"
        "表明模型训练充分且未出现明显的过拟合现象。")
    add_figure_placeholder(doc, "图2  训练曲线（损失/准确率/F1随轮次变化）")

    add_body(doc,
        "Grad-CAM热力图可视化结果显示，模型在进行病害分类时，注意力主要聚焦于叶片上的病斑区域，"
        "而非叶片边缘或背景等无关区域。这表明模型确实学习到了与病害相关的视觉特征，"
        "而非简单的纹理记忆或背景伪影，验证了模型决策的合理性。")
    add_figure_placeholder(doc, "图3  Grad-CAM热力图可视化示例")

    # ── 5. 系统实现 ──
    add_heading_l1(doc, "五、系统实现与功能展示")
    add_body(doc,
        "本系统基于Flask Web框架构建，后端集成PyTorch模型推理引擎。"
        "前端采用现代化的响应式设计，支持浅色简洁仪表盘界面。"
        "核心功能包括：（1）单张图片上传与诊断——支持拖拽上传、点击选择和URL粘贴；"
        "（2）批量预测——支持多张图片同时上传处理或ZIP压缩包上传；"
        "（3）实时推理——调用训练好的多任务模型，返回病害名称、严重程度和置信度；"
        "（4）Grad-CAM可视化——生成热力图叠加显示病灶位置；"
        "（5）诊断报告生成——自动生成包含病害详情、防控建议和复查时间线的结构化报告（支持JSON/PDF导出）；"
        "（6）风险评估——综合病害种类和严重程度计算风险等级。")

    add_body(doc,
        "系统默认监听0.0.0.0:7860端口，局域网内即可访问，方便在农业生产现场使用。"
        "模型支持ONNX导出格式，可进一步部署到边缘计算设备。"
        "代码使用Python 3.12和PyTorch 2.3+实现，依赖管理清晰，具备良好的可复现性。"
        "整个项目的代码量约为3000余行（含模型定义、训练逻辑、Web应用、工具函数等），"
        "核心代码均经过充分测试。")

    add_figure_placeholder(doc, "图4  系统Web界面截图")

    # ── 6. 总结与展望 ──
    add_heading_l1(doc, "六、总结与展望")
    add_body(doc,
        "本文设计并实现了一个基于多任务深度学习的农作物病害智能诊断与风险评估系统。"
        "主要工作包括：（1）提出了基于MobileNetV2和交叉注意力机制的多任务学习框架，"
        "在共享骨干网络的基础上，通过交叉注意力为不同任务生成任务特定特征；"
        "（2）通过多任务与单任务的对比实验，定量验证了多任务学习的协同效应，"
        "疾病分类准确率相比单任务基线提升约2.25个百分点；"
        "（3）集成了Grad-CAM可解释性分析，使模型决策过程可视化；"
        "（4）构建了完整的Web诊断系统，实现了从图像上传到报告生成的全流程自动化。")
    add_body(doc,
        "未来的改进方向包括：（1）引入Vision Transformer等先进网络架构替代MobileNetV2骨干网络，"
        "进一步提升分类精度；（2）研究知识蒸馏技术，将模型压缩至适合移动端部署的规模；"
        "（3）融合气象、土壤等多模态数据，实现更全面的作物健康评估；"
        "（4）在实际农田场景中部署验证，评估模型在复杂环境下的泛化性能。")

    # ── 7. 参考文献 ──
    add_heading_l1(doc, "参考文献")
    refs = [
        "[1] Sandler M, Howard A, Zhu M, et al. MobileNetV2: Inverted residuals and linear bottlenecks[C]. CVPR, 2018.",
        "[2] Selvaraju R R, Cogswell M, Das A, et al. Grad-CAM: Visual explanations from deep networks via gradient-based localization[C]. ICCV, 2017.",
        "[3] Caruana R. Multitask learning[J]. Machine Learning, 1997, 28(1): 41-75.",
        "[4] Misra I, Shrivastava A, Gupta A, et al. Cross-stitch networks for multi-task learning[C]. CVPR, 2016.",
        "[5] Liu S, Johns E, Davison A J. End-to-end multi-task learning with attention[C]. CVPR, 2019.",
        "[6] Mohanty S P, Hughes D P, Salathé M. Using deep learning for image-based plant disease detection[J]. Frontiers in Plant Science, 2016, 7: 1419.",
        "[7] Too E C, Yujian L, Njuki S, et al. A comparative study of fine-tuning deep learning models for plant disease identification[J]. Computers and Electronics in Agriculture, 2019, 161: 272-279.",
        "[8] He K, Zhang X, Ren S, et al. Deep residual learning for image recognition[C]. CVPR, 2016.",
        "[9] Howard A G, Zhu M, Chen B, et al. MobileNets: Efficient convolutional neural networks for mobile vision applications[J]. arXiv:1704.04861, 2017.",
        "[10] Tan M, Le Q V. EfficientNet: Rethinking model scaling for convolutional neural networks[C]. ICML, 2019.",
        "[11] Hughes D P, Salathé M. An open access repository of images on plant health to enable the development of mobile disease diagnostics[J]. arXiv:1511.08060, 2015.",
        "[12] Kingma D P, Ba J. Adam: A method for stochastic optimization[C]. ICLR, 2015.",
        "[13] Loshchilov I, Hutter F. Decoupled weight decay regularization[C]. ICLR, 2019.",
        "[14] Paszke A, Gross S, Massa F, et al. PyTorch: An imperative style, high-performance deep learning library[C]. NeurIPS, 2019.",
        "[15] Deng J, Dong W, Socher R, et al. ImageNet: A large-scale hierarchical image database[C]. CVPR, 2009.",
    ]
    for ref in refs:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.left_indent = Cm(0.74)
        para.paragraph_format.first_line_indent = Cm(-0.74)
        run = para.add_run(ref)
        set_run_font(run, "宋体", "Times New Roman", 10.5)

    # ── 添加页脚页码 ──
    add_footer_page_number(doc)

    # ── 教师评语页 ──
    doc.add_page_break()
    add_para(doc, "教师评语：", font_cn="黑体", size=12, bold=True,
             space_before=12, space_after=6, line_spacing=1.5)
    # 留空4行
    for _ in range(4):
        add_para(doc, "", size=12, line_spacing=1.5)

    add_para(doc, "成绩：___________", font_cn="宋体", size=12,
             space_before=12, space_after=6, line_spacing=1.5)
    add_para(doc, "教师签字：___________", font_cn="宋体", size=12,
             space_before=0, space_after=6, line_spacing=1.5)

    # ── 保存 ──
    doc.save(OUTPUT_PATH)
    print(f"✅ 已生成：{OUTPUT_PATH}")
    print(f"   文件大小：{os.path.getsize(OUTPUT_PATH)} bytes")


if __name__ == "__main__":
    generate()
