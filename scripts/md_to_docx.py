"""
Markdown 论文 → Word (.docx) 转换脚本
适用于《基于多任务深度学习的农作物病害智能诊断系统》论文
生成符合 数维杯 格式的 Word 文档
"""

import re
import os
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============ 配置 ============
MD_PATH = Path(__file__).resolve().parent.parent / "docs" / "论文_基于多任务深度学习的农作物病害智能诊断系统.md"
OUTPUT_DIR = Path(__file__).resolve().parent.parent / "docs"
OUTPUT_NAME = "论文_基于多任务深度学习的农作物病害智能诊断系统.docx"
IMG_BASE = Path(__file__).resolve().parent.parent / "docs" / "img"

TEAM_NUMBER = "SZUCM2026001"
TOTAL_PAGES = 120  # 占位符，生成后手动调整

# ============ 辅助函数 ============

def set_cell_shading(cell, color):
    """设置表格单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name_cn="宋体", name_en="Times New Roman", size=12, bold=False, italic=False, color=None):
    """设置 run 的字体属性"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = name_en
    # 设置中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph_with_format(doc, text, style=None, alignment=None, font_name_cn="宋体",
                               font_name_en="Times New Roman", font_size=12, bold=False,
                               space_before=0, space_after=6, first_line_indent=None,
                               color=None, italic=False):
    """添加带格式的段落"""
    p = doc.add_paragraph()
    if style:
        p.style = style
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)

    run = p.add_run(text)
    set_run_font(run, font_name_cn, font_name_en, font_size, bold, italic, color)
    return p


def add_page_number(doc):
    """添加页码"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Team# 前缀
        run = p.add_run(f"Team#{TEAM_NUMBER}    ")
        set_run_font(run, "宋体", "Times New Roman", 9)

        # 页码域
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)

        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)

        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)

        # 总页数
        run4 = p.add_run(f"  of {TOTAL_PAGES}")
        set_run_font(run4, "宋体", "Times New Roman", 9)


def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), "黑体")
    return h


def add_code_block(doc, code_text):
    """添加代码块（等宽字体、灰色背景）"""
    for line in code_text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1.0)
        # 灰色底纹
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
        pPr = p._element.get_or_add_pPr()
        pPr.append(shading)

        run = p.add_run(line if line else " ")
        set_run_font(run, "Courier New", "Courier New", 8)


def add_image_with_caption(doc, img_path, caption, img_width=5.5):
    """添加图片及图注"""
    if not os.path.exists(img_path):
        add_paragraph_with_format(doc, f"[图片未找到: {img_path}]",
                                  font_size=10, italic=True, color=(200, 0, 0))
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(img_path, width=Inches(img_width))

    # 图注
    add_paragraph_with_format(doc, caption, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              font_size=10, bold=False, italic=True, space_before=2, space_after=10)


def add_table_from_data(doc, headers, rows, caption=None):
    """添加表格"""
    if caption:
        add_paragraph_with_format(doc, caption, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  font_size=10, bold=True, space_before=6, space_after=4)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, "黑体", "Times New Roman", 10, bold=True)
        set_cell_shading(cell, "D9E2F3")

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            set_run_font(run, "宋体", "Times New Roman", 9)

    if caption:
        doc.add_paragraph()  # spacing


# ============ Markdown 解析与转换 ============

class MarkdownToDocx:
    def __init__(self, md_path, output_path):
        self.md_path = md_path
        self.output_path = output_path
        self.doc = Document()

        # 页面设置
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)   # A4
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        self.in_code_block = False
        self.code_buffer = []
        self.in_table = False
        self.table_lines = []

    def parse_inline_formatting(self, text):
        """处理行内格式：加粗、斜体、行内代码、数学公式"""
        parts = []
        last_end = 0

        # 匹配各种行内格式
        patterns = [
            (r'\$\$(.*?)\$\$', 'math_block'),
            (r'\$(.*?)\$', 'math_inline'),
            (r'`([^`]+)`', 'code'),
            (r'\*\*(.+?)\*\*', 'bold'),
            (r'\*(.+?)\*', 'italic'),
        ]

        # 简单版本：逐个字符处理
        # 采用正则依次匹配
        pos = 0
        while pos < len(text):
            best_match = None
            best_type = None
            best_end = None

            for pattern, ptype in patterns:
                m = re.search(pattern, text[pos:])
                if m:
                    if best_match is None or m.start() < best_match.start():
                        best_match = m
                        best_type = ptype
                        best_end = m.end()

            if best_match is None:
                parts.append(('text', text[pos:]))
                break

            # 匹配前的普通文本
            if best_match.start() > 0:
                parts.append(('text', text[pos:pos + best_match.start()]))

            # 匹配的格式文本
            content = best_match.group(1) if len(best_match.groups()) >= 1 else best_match.group(0)
            parts.append((best_type, content))
            pos += best_end

        return parts

    def add_rich_paragraph(self, text, **kwargs):
        """添加富文本段落（支持行内格式）"""
        p = self.doc.add_paragraph()
        alignment = kwargs.get('alignment')
        if alignment is not None:
            p.alignment = alignment
        p.paragraph_format.space_before = Pt(kwargs.get('space_before', 3))
        p.paragraph_format.space_after = Pt(kwargs.get('space_after', 3))
        first_line_indent = kwargs.get('first_line_indent')
        if first_line_indent:
            p.paragraph_format.first_line_indent = Cm(first_line_indent)

        default_size = kwargs.get('font_size', 12)

        parts = self.parse_inline_formatting(text)
        for ptype, content in parts:
            run = p.add_run()
            if ptype == 'text':
                set_run_font(run, "宋体", "Times New Roman", default_size)
                run.text = content
            elif ptype == 'bold':
                set_run_font(run, "黑体", "Times New Roman", default_size, bold=True)
                run.text = content
            elif ptype == 'italic':
                set_run_font(run, "宋体", "Times New Roman", default_size, italic=True)
                run.text = content
            elif ptype == 'code':
                set_run_font(run, "Courier New", "Courier New", default_size - 1)
                run.text = content
            elif ptype == 'math_inline':
                set_run_font(run, "Cambria Math", "Cambria Math", default_size, italic=True)
                run.text = content
            elif ptype == 'math_block':
                set_run_font(run, "Cambria Math", "Cambria Math", default_size, italic=True)
                run.text = content

        return p

    def _flush_table(self):
        """刷新并输出缓存的表格"""
        if len(self.table_lines) < 2:
            self.table_lines = []
            return
        headers = self.table_lines[0]
        rows = self.table_lines[1:]
        max_cols = max(len(r) for r in rows)
        while len(headers) < max_cols:
            headers.append("")
        for i, row in enumerate(rows):
            while len(row) < max_cols:
                row.append("")
        add_table_from_data(self.doc, headers, rows)
        self.table_lines = []

    def process_line(self, line):
        """处理单行 Markdown"""
        stripped = line.strip()

        # 空行
        if not stripped:
            if self.in_code_block:
                self.code_buffer.append("")
            return

        # 代码块
        if stripped.startswith("```"):
            if self.in_code_block:
                code_text = "\n".join(self.code_buffer)
                add_code_block(self.doc, code_text)
                self.code_buffer = []
                self.in_code_block = False
            else:
                self.in_code_block = True
                self.code_buffer = []
            return

        if self.in_code_block:
            self.code_buffer.append(line)
            return

        # 水平线
        if stripped.startswith("---") or stripped.startswith("***") or stripped.startswith("___"):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            return

        # 表格行（检测 | 分隔的表格）
        if "|" in stripped and stripped.count("|") >= 3:
            if re.match(r'^[\s\|:\-]+$', stripped):
                return
            cells = [c.strip() for c in stripped.split("|") if c.strip()]
            self.table_lines.append(cells)
            return

        # 非表格行 — 刷新表格缓存
        self._flush_table()

        # 标题 (#)
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            add_heading_custom(self.doc, text, level=level)
            return

        # 图片 ![]()
        img_match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_rel_path = img_match.group(2)
            img_path = img_rel_path
            if not os.path.isabs(img_path):
                candidates = [
                    Path(img_rel_path).resolve(),
                    IMG_BASE / Path(img_rel_path).name,
                    Path(__file__).resolve().parent.parent / img_rel_path,
                    Path(__file__).resolve().parent.parent / "docs" / img_rel_path,
                ]
                for c in candidates:
                    if c.exists():
                        img_path = str(c)
                        break
            width_match = re.search(r'\{:\s*width\s*=\s*(\d+)%\s*\}', stripped)
            img_width = 5.5
            if width_match:
                pct = int(width_match.group(1)) / 100
                img_width = 5.5 * pct
            add_image_with_caption(self.doc, img_path, alt_text, img_width=img_width)
            return

        # 列表项
        list_match = re.match(r'^(\s*)[\-\*]\s+(.+)$', stripped)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(2)
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0 + indent * 0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run("• " + text)
            set_run_font(run, "宋体", "Times New Roman", 12)
            return

        # 编号列表
        ol_match = re.match(r'^\s*(\d+)\.\s+(.+)$', stripped)
        if ol_match:
            num = ol_match.group(1)
            text = ol_match.group(2)
            self.add_rich_paragraph(f"  {num}. {text}", first_line_indent=0)
            return

        # 普通段落
        self.add_rich_paragraph(stripped, first_line_indent=0.74)

    def convert(self):
        """执行转换"""
        print(f"正在读取: {self.md_path}")

        with open(self.md_path, "r", encoding="utf-8") as f:
            lines = f.readlines()

        print(f"共 {len(lines)} 行，开始转换...")

        # 处理首页 Team# 页眉
        # 读取第一行获取 Team# 信息
        first_line = lines[0].strip() if lines else ""
        team_match = re.match(r'Team#(\S+)', first_line)

        # 添加标题页
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(72)
        run = p.add_run("基于多任务深度学习的\n农作物病害智能诊断系统")
        set_run_font(run, "黑体", "Times New Roman", 26, bold=True)

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(24)
        run = p2.add_run(f"Team#{TEAM_NUMBER}")
        set_run_font(run, "Times New Roman", "Times New Roman", 16, bold=True)

        p3 = self.doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = Pt(12)
        run = p3.add_run("机器学习课程大作业")
        set_run_font(run, "宋体", "Times New Roman", 14)

        self.doc.add_page_break()

        # 处理正文
        for line in lines[1:]:  # 跳过第一行（Team# header）
            self.process_line(line)
        # 结束前刷新最后一张表格
        self._flush_table()

        # 添加页码
        add_page_number(self.doc)

        # 保存
        self.doc.save(str(self.output_path))
        print(f"✓ 成功生成: {self.output_path}")
        print(f"  文件大小: {os.path.getsize(str(self.output_path)) / 1024:.1f} KB")


if __name__ == "__main__":
    # 切换到项目根目录
    os.chdir(Path(__file__).resolve().parent.parent)

    output_path = OUTPUT_DIR / OUTPUT_NAME
    converter = MarkdownToDocx(MD_PATH, output_path)
    converter.convert()
